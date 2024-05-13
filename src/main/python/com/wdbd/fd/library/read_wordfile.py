#!/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File    :   read_wordfile.py
@Time    :   2024/05/13 14:18:45
@Author  :   Jeffrey Wang
@Version :   1.0
@Contact :   shwangjj@163.com
@Desc    :   尝试读取Word文件中大纲数据
'''
import os
from docx import Document

def read_all_sytles(doc_path):
    """ 打印所有样式名称 """
    document = Document(doc_path)
    # print(document.styles)
    
    for i, paragraph in enumerate(document.paragraphs):
        print(f"sytle = {paragraph.style.name}")
    
    # # 有Sytle的段落：
    # for style in document.styles:
    #     # 打印样式名称和类型
    #     print(f"Style Name: {style.name}, Style Type: {style.type}")
    
    # 无Style的段落：
    # for idx, p in enumerate(document.paragraphs):
    #     print(f"第{idx}段, 内容：{p.text}")
    #     # print(p.style)
    #     xml_str = p._p.xml
    #     print(xml_str)
    #     print("sytle = {}".format(get_sytle_name(xml_str)))

def get_sytle_name(xml_string):
    """ 取得段落样式名称 """
    t_str = "<w:pStyle w:val=\""
    start = xml_string.find(t_str) + len(t_str)
    # print(start)
    end = xml_string.find("\"", start+1)
    # print(end)
    # print(xml_string[start: end])
    return xml_string[start: end]

def read_word_outline(doc_path):
    """ 读取文件大纲 """
    document = Document(doc_path)
    outline_levels = []
    for i, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.style:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name[7:]
                outline_levels.append((i, level, paragraph.text))
            else:
                # 非windows操作系统，没有Heading1-9的样式
                return [(0, -1, "无法识别的格式")]
    if len(outline_levels) == 0:
        return [(0, -1, "没有找到大纲")]
    return outline_levels

def scan_directory_for_docx(directory):
    """ 扫描目录，返回docx文件列表 """
    
    docx_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx'):
                docx_files.append(os.path.abspath(os.path.join(root, file)))
    return docx_files


def scan_and_parse(directory):
    """ 扫描目录，返回docx文件列表 """
    if not os.path.exists(directory):
        print(f"文件夹路径错误，请检查! 路径：{directory}")
        return
    
    outlines = []
    docx_file_paths = scan_directory_for_docx(directory)
    # print(f"一共有{len(docx_file_paths)}个文件")
    outlines.append(f"一共有{len(docx_file_paths)}个文件")
    for file_path in docx_file_paths:
        # print(f"文件：{file_path}")
        outlines.append(f"文件：{file_path}")
        outline = read_word_outline(file_path)
        for idx, level, text in outline:
            output = "{idx}段，{tab}级别:{level}, 标题：{text}".format(idx=idx, level=level, text=text, tab=(int(level)*'  '))
            outlines.append(output)
        outlines.append("\n")
    return outlines    

if __name__ == "__main__":
    # 解析特定目录下的所有docx文件，并将大纲输出到文件
    directory_to_scan = '/Users/junjiewang/python_prj/projects/feedme'  # 替换为你的目录路径
    to_file = "scan.txt"    # 替换为输出文件的路径

    outputs = scan_and_parse(directory_to_scan)
    # 输出命令行
    for line in outputs:
        print(line)
    # 输出到文件
    with open(to_file, 'w', encoding='utf-8') as file:
        # 遍历列表中的每个字符串
        for string in outputs:
            # 将字符串写入文件，每个字符串后添加一个换行符
            file.write(string + '\n')
